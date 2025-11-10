import os
import re
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import json
import PyPDF2
from docx import Document
from groq import Groq
from collections import defaultdict
import numpy as np
from env import GROQ_API_KEY
import sys

class ResumeParser:
    def __init__(self, api_key: str):
        """Initialize the resume parser with Groq API key."""
        # Backward compatible Groq client initialization
        try:
            # Try modern initialization first
            self.client = Groq(api_key=api_key)
        except TypeError as e:
            if "proxies" in str(e) or "unexpected keyword argument" in str(e):
                # Fallback for older versions - initialize with minimal params
                try:
                    # Try with just api_key
                    import groq
                    
                    # Check if we need to use legacy initialization
                    if hasattr(groq, '__version__'):
                        version = groq.__version__
                        print(f"Detected Groq version: {version}")
                    
                    # Attempt basic initialization without optional parameters
                    self.client = Groq(api_key=api_key)
                    
                except Exception as fallback_error:
                    print(f"Warning: Groq initialization fallback failed: {fallback_error}")
                    # Last resort: try to create client with absolute minimum
                    try:
                        from groq import Groq as GroqClient
                        self.client = GroqClient(api_key=api_key)
                    except:
                        raise RuntimeError(
                            f"Unable to initialize Groq client. "
                            f"Please ensure groq>=0.4.0 is installed: pip install --upgrade groq"
                        ) from e
            else:
                raise
        
    def extract_text_from_pdf(self, pdf_path: str) -> Tuple[str, str]:
        """Extract text from PDF with advanced bi-column handling.
        Returns: (formatted_text, layout_metadata)
        """
        try:
            import pdfplumber
            all_text_blocks = []
            layout_info = []
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract words with complete position information
                    words = page.extract_words(
                        x_tolerance=2,
                        y_tolerance=2,
                        keep_blank_chars=False,
                        use_text_flow=False,
                        extra_attrs=['size', 'fontname']
                    )
                    
                    if not words:
                        continue
                    
                    # Detect layout structure
                    layout = self._analyze_page_layout(words, page.width, page.height)
                    layout_info.append({
                        'page': page_num + 1,
                        'layout_type': layout['type'],
                        'columns': layout['columns']
                    })
                    
                    # Extract text based on layout
                    page_text = self._extract_structured_text(words, layout, page.width, page.height)
                    all_text_blocks.append(page_text)
            
            formatted_text = "\n\n=== PAGE BREAK ===\n\n".join(all_text_blocks)
            metadata = json.dumps(layout_info, indent=2)
            
            return formatted_text, metadata
            
        except ImportError:
            # Fallback to PyPDF2
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text, "{}"
    
    def _analyze_page_layout(self, words: List[Dict], page_width: float, page_height: float) -> Dict:
        """Analyze page layout to detect columns and structure."""
        if not words:
            return {'type': 'single_column', 'columns': [(0, page_width)]}
        
        # Get x-positions of all word centers
        x_positions = [(w['x0'] + w['x1']) / 2 for w in words]
        
        # Enhanced column detection with better sensitivity for 2-column layouts
        if len(x_positions) > 10:
            try:
                from sklearn.cluster import KMeans
            except ImportError:
                # Fallback if sklearn not available
                return {'type': 'single_column', 'columns': [(0, page_width)]}
            
            X = np.array(x_positions).reshape(-1, 1)
            
            # Calculate horizontal spread to detect columns more reliably
            x_spread = np.std(x_positions)
            page_center = page_width / 2
            
            # Count words on left vs right half
            left_count = sum(1 for x in x_positions if x < page_center)
            right_count = len(x_positions) - left_count
            balance_ratio = min(left_count, right_count) / max(left_count, right_count) if max(left_count, right_count) > 0 else 0
            
            # Try 1, 2, and 3 column layouts with adjusted scoring
            best_score = float('inf')
            best_k = 1
            
            for k in [1, 2, 3]:
                if len(x_positions) < k * 5:  # Need enough samples
                    continue
                    
                kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
                kmeans.fit(X)
                score = kmeans.inertia_
                
                # Enhanced scoring for 2-column detection
                if k == 2:
                    # Favor 2-column if there's good balance and spread
                    if balance_ratio > 0.3 and x_spread > page_width * 0.15:
                        score *= 0.5  # Strong preference for 2-column
                    elif balance_ratio > 0.2:
                        score *= 0.7
                elif k > 2:
                    score *= (1 + 0.4 * k)  # Higher penalty for over-clustering
                
                if score < best_score:
                    best_score = score
                    best_k = k
            
            if best_k > 1:
                # Multi-column layout detected
                kmeans = KMeans(n_clusters=best_k, random_state=42, n_init=10)
                kmeans.fit(X)
                centers = sorted(kmeans.cluster_centers_.flatten().tolist())
                
                # Calculate column boundaries with better gap detection
                columns = []
                margin = page_width * 0.05
                
                for i, center in enumerate(centers):
                    if i == 0:
                        col_start = 0
                    else:
                        col_start = (centers[i-1] + center) / 2
                    
                    if i == len(centers) - 1:
                        col_end = page_width
                    else:
                        col_end = (center + centers[i+1]) / 2
                    
                    columns.append((col_start, col_end))
                
                return {
                    'type': f'{best_k}_column',
                    'columns': columns,
                    'column_centers': centers,
                    'balance_ratio': balance_ratio
                }
        
        # Default to single column
        return {'type': 'single_column', 'columns': [(0, page_width)]}
    
    def _extract_structured_text(self, words: List[Dict], layout: Dict, page_width: float, page_height: float) -> str:
        """Extract text preserving structure based on layout."""
        # Group words into text blocks
        text_blocks = self._group_into_blocks(words)
        
        # Classify blocks (header, section title, body text)
        classified_blocks = self._classify_blocks(text_blocks, page_width)
        
        # Sort blocks by reading order
        sorted_blocks = self._sort_by_reading_order(classified_blocks, layout)
        
        # Format output with structure markers and column indicators
        formatted_lines = []
        current_column = 0
        
        # Add layout type indicator for multi-column resumes
        if layout['type'] != 'single_column':
            formatted_lines.append(f"[LAYOUT: {layout['type'].upper()}]")
        
        for i, block in enumerate(sorted_blocks):
            block_type = block['type']
            text = block['text'].strip()
            
            if not text:
                continue
            
            # Detect column transitions in 2-column layout
            if layout['type'] == '2_column' and i > 0:
                block_center = (block['x0'] + block['x1']) / 2
                prev_center = (sorted_blocks[i-1]['x0'] + sorted_blocks[i-1]['x1']) / 2
                
                # Check if we've moved to a different column
                new_column = 0 if block_center < page_width / 2 else 1
                prev_column = 0 if prev_center < page_width / 2 else 1
                
                if new_column != prev_column:
                    formatted_lines.append(f"\n[COLUMN {new_column + 1}]")
                    current_column = new_column
            
            if block_type == 'section_header':
                formatted_lines.append(f"\n### {text} ###")
            elif block_type == 'subsection_header':
                formatted_lines.append(f"\n## {text} ##")
            elif block_type == 'full_width_header':
                formatted_lines.append(f"\n# {text} #")
            else:
                formatted_lines.append(text)
        
        return '\n'.join(formatted_lines)
    
    def _group_into_blocks(self, words: List[Dict]) -> List[Dict]:
        """Group words into logical text blocks."""
        if not words:
            return []
        
        blocks = []
        current_block = {
            'words': [words[0]],
            'x0': words[0]['x0'],
            'y0': words[0]['top'],
            'x1': words[0]['x1'],
            'y1': words[0]['bottom'],
            'size': words[0].get('size', 12),
            'fontname': words[0].get('fontname', '')
        }
        
        y_tolerance = 3
        x_tolerance = 5
        
        for word in words[1:]:
            # Check if word belongs to current block
            same_line = abs(word['top'] - current_block['y0']) < y_tolerance
            continues_horizontally = (word['x0'] - current_block['x1']) < x_tolerance
            same_font = word.get('fontname', '') == current_block['fontname']
            
            if same_line and continues_horizontally:
                # Add to current block
                current_block['words'].append(word)
                current_block['x1'] = word['x1']
                current_block['y1'] = max(current_block['y1'], word['bottom'])
            else:
                # Start new block
                blocks.append(current_block)
                current_block = {
                    'words': [word],
                    'x0': word['x0'],
                    'y0': word['top'],
                    'x1': word['x1'],
                    'y1': word['bottom'],
                    'size': word.get('size', 12),
                    'fontname': word.get('fontname', '')
                }
        
        blocks.append(current_block)
        
        # Merge adjacent blocks on same line
        merged_blocks = []
        for block in blocks:
            block['text'] = ' '.join(w['text'] for w in block['words'])
            merged_blocks.append(block)
        
        return merged_blocks
    
    def _classify_blocks(self, blocks: List[Dict], page_width: float) -> List[Dict]:
        """Classify text blocks by type (header, body, etc)."""
        if not blocks:
            return []
        
        # Calculate average font size
        avg_size = sum(b['size'] for b in blocks) / len(blocks)
        
        for block in blocks:
            text = block['text']
            size = block['size']
            width = block['x1'] - block['x0']
            
            # Determine block type
            is_large = size > avg_size * 1.2
            is_bold = 'bold' in block['fontname'].lower()
            is_all_caps = text.isupper() and len(text) > 3
            is_full_width = width > page_width * 0.7
            
            if is_full_width and (is_large or is_all_caps or is_bold):
                block['type'] = 'full_width_header'
            elif (is_large or is_all_caps or is_bold) and len(text) < 100:
                block['type'] = 'section_header'
            elif is_bold and len(text) < 60:
                block['type'] = 'subsection_header'
            else:
                block['type'] = 'body_text'
        
        return blocks
    
    def _sort_by_reading_order(self, blocks: List[Dict], layout: Dict) -> List[Dict]:
        """Sort blocks in correct reading order for bi-column layouts."""
        if layout['type'] == 'single_column':
            # Simple top-to-bottom sort
            return sorted(blocks, key=lambda b: (b['y0'], b['x0']))
        
        # Multi-column layout - enhanced for better section capture
        columns = layout['columns']
        
        # Separate full-width headers from column content
        full_width_blocks = [b for b in blocks if b['type'] == 'full_width_header']
        column_blocks = [b for b in blocks if b['type'] != 'full_width_header']
        
        # Group blocks by vertical zones with reduced threshold for better section detection
        zones = self._group_into_zones(column_blocks, zone_threshold=25)
        
        sorted_blocks = []
        
        for zone_blocks in zones:
            # Check if this zone has full-width header
            zone_y = min(b['y0'] for b in zone_blocks)
            zone_headers = [b for b in full_width_blocks if abs(b['y0'] - zone_y) < 20]
            sorted_blocks.extend(sorted(zone_headers, key=lambda b: b['y0']))
            
            # Sort blocks within zone by column, then by vertical position
            for col_start, col_end in columns:
                col_blocks = [
                    b for b in zone_blocks
                    if col_start <= ((b['x0'] + b['x1']) / 2) < col_end
                ]
                # Sort by y-position within each column
                sorted_blocks.extend(sorted(col_blocks, key=lambda b: (b['y0'], b['x0'])))
        
        return sorted_blocks
    
    def _group_into_zones(self, blocks: List[Dict], zone_threshold: int = 30) -> List[List[Dict]]:
        """Group blocks into vertical zones with adjustable threshold."""
        if not blocks:
            return []
        
        # Sort by y-position
        sorted_blocks = sorted(blocks, key=lambda b: b['y0'])
        
        zones = []
        current_zone = [sorted_blocks[0]]
        
        for block in sorted_blocks[1:]:
            # Check if block is in same zone as previous
            prev_bottom = max(b['y1'] for b in current_zone)
            
            if block['y0'] - prev_bottom < zone_threshold:
                current_zone.append(block)
            else:
                zones.append(current_zone)
                current_zone = [block]
        
        zones.append(current_zone)
        return zones
    
    def extract_text_from_docx(self, docx_path: str) -> Tuple[str, str]:
        """Extract text from DOCX file."""
        doc = Document(docx_path)
        
        formatted_lines = []
        
        # Extract from paragraphs with style information
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headers by style
            style_name = para.style.name.lower()
            
            if 'heading 1' in style_name or para.runs and para.runs[0].bold and len(text) < 100:
                formatted_lines.append(f"\n### {text} ###")
            elif 'heading' in style_name or (para.runs and para.runs[0].bold and len(text) < 60):
                formatted_lines.append(f"\n## {text} ##")
            else:
                formatted_lines.append(text)
        
        # Extract from tables
        for table in doc.tables:
            formatted_lines.append("\n[TABLE]")
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    formatted_lines.append(row_text)
            formatted_lines.append("[/TABLE]\n")
        
        formatted_text = '\n'.join(formatted_lines)
        metadata = json.dumps({'source': 'docx', 'layout_type': 'document'})
        
        return formatted_text, metadata
    
    def _is_valid_email(self, text: str) -> bool:
        """Validate if text is a proper email address."""
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        
        if not re.match(email_pattern, text.strip()):
            return False
        
        parts = text.strip().split('@')
        if len(parts) != 2:
            return False
        
        local, domain = parts
        
        if not local or len(local) > 64:
            return False
        
        if not domain or '.' not in domain:
            return False
        
        tld = domain.split('.')[-1]
        if len(tld) < 2 or not tld.isalpha():
            return False
        
        return True
    
    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information with improved validation."""
        contact_info = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None,
            'location': None
        }
        
        # Extract email with validation
        email_matches = re.findall(r'\S+@\S+\.\S+', text)
        for email in email_matches:
            # Clean up email
            email = re.sub(r'[,;:)\]]$', '', email)
            if self._is_valid_email(email):
                contact_info['email'] = email.strip()
                break
        
        # Extract phone with multiple formats
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{3}[-.\s]\d{3}[-.\s]\d{4}',
            r'\+?\d{10,15}'
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact_info['phone'] = phone_match.group(0).strip()
                break
        
        # Extract LinkedIn
        linkedin_patterns = [
            r'linkedin\.com/in/[\w-]+',
            r'linkedin\.com/[\w-]+'
        ]
        for pattern in linkedin_patterns:
            linkedin_match = re.search(pattern, text, re.IGNORECASE)
            if linkedin_match:
                contact_info['linkedin'] = linkedin_match.group(0)
                break
        
        # Extract GitHub
        github_patterns = [
            r'github\.com/[\w-]+',
        ]
        for pattern in github_patterns:
            github_match = re.search(pattern, text, re.IGNORECASE)
            if github_match:
                contact_info['github'] = github_match.group(0)
                break
        
        return contact_info
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Parse resume with advanced bi-column support and dynamic sections."""
        # Extract text based on file type
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            formatted_text, layout_metadata = self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            formatted_text, layout_metadata = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Extract basic contact info
        contact_info = self.extract_contact_info(formatted_text)
        
        # Enhanced prompt for better 2-column handling
        prompt = f"""Parse the following resume text and extract ALL information comprehensively.

CRITICAL INSTRUCTIONS FOR 2-COLUMN RESUMES:
1. This resume may have a 2-COLUMN LAYOUT - pay special attention to [COLUMN 1] and [COLUMN 2] markers
2. Text marked with ### indicates section headers - these are CRITICAL for identifying all sections
3. DO NOT miss any sections - both columns contain important information
4. Common 2-column pattern: Left column (skills, education, certifications) | Right column (experience, projects)
5. Extract EVERY section found, regardless of which column it's in
6. DO NOT use predefined section names - use EXACT names from the resume
7. For emails: Only extract valid email addresses (must have proper domain like @gmail.com, @company.com)
8. Preserve the hierarchical structure of content within each column
9. Extract links from all hyperlinked text or image in the Resume. Analyse all added images for links.

Layout metadata: {layout_metadata}

Resume Text:
{formatted_text}

Return a JSON object with this structure:
{{
  "personal_info": {{
    "name": "extracted name",
    "email": "valid email only",
    "phone": "phone number",
    "location": "city, state/country",
    "linkedin": "url",
    "github": "url",
    "portfolio": "url",
    "other_links": ["any other professional links"]
  }},
  "sections": [
    {{
      "section_name": "EXACT name as it appears in resume (e.g., 'TECHNICAL SKILLS', 'WORK EXPERIENCE', 'CERTIFICATIONS')",
      "content_type": "summary|experience|education|skills|projects|certifications|awards|publications|languages|volunteering|custom",
      "column": "left|right|full_width",
      "items": [
        // For experience/education/projects: array of objects with all relevant fields (title, company, dates, description, etc.)
        // For skills: array of strings or categorized objects
        // For summary: single string or array of key points
        // For certifications/awards: array with name, issuer, date
        // For custom sections: preserve original structure exactly
      ]
    }}
  ]
}}

SECTION EXTRACTION CHECKLIST - Ensure you capture ALL of these if present:
- Professional Summary/Objective
- Work Experience/Professional Experience
- Education/Academic Background
- Technical Skills/Skills/Core Competencies
- Projects/Key Projects
- Certifications/Licenses
- Awards/Honors/Achievements
- Publications/Research
- Languages
- Volunteer Experience
- Professional Affiliations
- Any other custom sections

Be exhaustive - extract EVERY piece of information from BOTH columns.
Preserve original formatting, bullet points, dates, and maintain accurate hierarchical structure and listing structure.
Do not skip any section, no matter how small or unconventional the section name is.
"""
        
        response = self.client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=8000
        )
        
        # Parse Groq's response
        parsed_data = self._parse_claude_response(response.choices[0].message.content)
        
        # Merge with extracted contact info
        if 'personal_info' in parsed_data:
            for key, value in contact_info.items():
                if value and not parsed_data['personal_info'].get(key):
                    parsed_data['personal_info'][key] = value
        
        # Add metadata
        parsed_data['metadata'] = {
            'file_path': file_path,
            'layout_info': json.loads(layout_metadata) if layout_metadata else {},
            'parsed_at': datetime.now().isoformat()
        }
        
        return parsed_data
    
    def _parse_claude_response(self, response_text: str) -> Dict[str, Any]:
        """Parse Claude's JSON response."""
        # Extract JSON from response - handle markdown code blocks
        json_pattern = r'```json\s*(.*?)\s*```|```\s*(.*?)\s*```|\{.*\}'
        json_match = re.search(json_pattern, response_text, re.DOTALL)
        
        if json_match:
            json_str = json_match.group(1) or json_match.group(2) or json_match.group(0)
            json_str = json_str.strip('`').strip()
            if json_str.startswith('json'):
                json_str = json_str[4:].strip()
            
            try:
                return json.loads(json_str)
            except json.JSONDecodeError as e:
                print(f"JSON decode error: {e}")
                print(f"Attempted to parse: {json_str[:200]}...")
        
        # Return structured default if parsing fails
        return {
            'personal_info': {},
            'sections': []
        }

# Example usage
if __name__ == "__main__":
    api_key = GROQ_API_KEY
    if not api_key:
        raise ValueError("GROQ_API_KEY environment variable not set")
    
    parser = ResumeParser(api_key)
    
    # Parse resume
    resume_data = parser.parse_resume("resume.pdf")
    
    print("Parsed Resume Data:")
    print(json.dumps(resume_data, indent=2))
    
    # Print section names found
    print("\nSections found in resume:")
    for section in resume_data.get('sections', []):
        column_info = f" [{section.get('column', 'unknown')}]" if 'column' in section else ""
        print(f"  - {section['section_name']} ({section['content_type']}){column_info}")
    
    with open("extracted_resume.json", 'w', encoding='utf-8') as f:
        json.dump(resume_data, f, indent=2, ensure_ascii=False)
